import docx


def person_invitation(wrd, name):
    wrd.add_heading("it would be a pleasure to have the company of", 2)
    wrd.add_heading(name, 1)
    wrd.add_heading("at 11010 memory lane on the evening of", 2)
    wrd.add_heading("at 7 o'clock", 2)
    para = wrd.add_paragraph(".")
    para.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)


if __name__ == "__main__":

    doc = docx.Document()

    with open("./guests.txt") as f:
        for name in f:
            person_invitation(doc, name)

    doc.save("invitations.docx")
